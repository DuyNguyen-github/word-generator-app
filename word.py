import os
import datetime
import re
import unicodedata
from openpyxl import load_workbook
from docx import Document

# --- Cấu hình ---
# Mặc định (có thể ghi đè khi gọi từ web)
DEFAULT_OUTPUT_FOLDER = "output"
os.makedirs(DEFAULT_OUTPUT_FOLDER, exist_ok=True)

# Danh sách cột mặc định (phải trùng EXACT với tên cột trong Excel và placeholder trong Word)
DEFAULT_COLUMNS_TO_USE = [
    "Họ và tên","Đơn vị công tác mới của Chi nhánh","Chức vụ, đơn vị đang công tác công việc hiện tại","Đơn vị công tác hiện tại",
    "Chức vụ, đơn vị đang công tác bố trí tại Chi nhánh"
]


def generate_from_files(excel_file, template_file, output_folder=None, columns_to_use=None):
    """Generate Word files for each row in the Excel file using the given template.

    Args:
        excel_file (str): path to the Excel file
        template_file (str): path to the .docx template with placeholders like {{ColumnName}}
        output_folder (str): folder where generated files are saved (created if missing)
        columns_to_use (list): list of column names to include in replacement (defaults to DEFAULT_COLUMNS_TO_USE)

    Returns:
        list: list of created output file paths
    """
    if output_folder is None:
        output_folder = DEFAULT_OUTPUT_FOLDER
    if columns_to_use is None:
        columns_to_use = DEFAULT_COLUMNS_TO_USE

    os.makedirs(output_folder, exist_ok=True)

    headers, rows = read_excel_display_rows(excel_file)
    # determine placeholders present in the template
    placeholders = get_placeholders_from_template(template_file)

    # build a mapping from placeholder -> matching header name in Excel (or None)
    header_norm_map = {h: normalize_str(h) for h in headers}
    placeholder_to_header = {}
    for ph in placeholders:
        # try exact match first
        if ph in headers:
            placeholder_to_header[ph] = ph
            continue

        # try normalized matching
        ph_norm = normalize_str(ph)
        matched = None
        for h, h_norm in header_norm_map.items():
            if h_norm == ph_norm:
                matched = h
                break
        placeholder_to_header[ph] = matched

    # log placeholders that weren't matched to any header
    unmatched = [p for p, h in placeholder_to_header.items() if h is None]
    if unmatched:
        print(f"Warning: the following placeholders were not matched to any Excel header and will be empty: {unmatched}")
    created = []

    for idx, row in enumerate(rows, start=1):
        # build data dict mapping placeholder -> replacement value (string)
        data = {}
        for ph, header_name in placeholder_to_header.items():
            if header_name is None:
                data[ph] = ""
            else:
                data[ph] = row.get(header_name, "")
        safe_name = re.sub(r'[\\/*?:"<>|]', "_", data.get("Họ và tên", f"row{idx}"))
        output_path = os.path.join(output_folder, f"{idx}_{safe_name}.docx")

        fill_template(template_file, output_path, data)
        created.append(output_path)

    return created


def normalize_str(s):
    """Normalize a string for fuzzy matching: lowercase, remove diacritics and non-alphanumerics."""
    if s is None:
        return ""
    # convert to str
    s = str(s)
    s = s.strip().lower()
    # remove diacritics
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(ch for ch in s if not unicodedata.combining(ch))
    # remove non-alphanumeric characters
    s = re.sub(r'[^0-9a-z]+', '', s)
    return s


def get_placeholders_from_template(template_path):
    """Extract placeholder names from a .docx template.

    Looks for patterns like {{placeholder}} across paragraphs and table cells. Returns a set of names (without braces).
    """
    doc = Document(template_path)
    pattern = re.compile(r"{{\s*([^{}]+?)\s*}}")
    found = set()

    # paragraphs
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        for m in pattern.findall(full):
            found.add(m)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(r.text for r in p.runs)
                    for m in pattern.findall(full):
                        found.add(m)

    return found

# --- Hàm đọc Excel bằng openpyxl và chuyển từng ô thành string hiển thị y chang Excel ---
def read_excel_display_rows(path):
    wb = load_workbook(path, data_only=True)
    ws = wb.active

    # Lấy header (người dùng đảm bảo header nằm ở hàng 1)
    headers = [cell.value for cell in ws[1]]
    rows = []
    for excel_row in ws.iter_rows(min_row=2, values_only=False):
        row_dict = {}
        for header, cell in zip(headers, excel_row):
            raw = cell.value
            if raw is None:
                display = ""
            elif isinstance(raw, (datetime.datetime, datetime.date)):
                # chuyển date/datetime sang dd/mm/YYYY
                display = raw.strftime("%d/%m/%Y")
            elif isinstance(raw, (int, float)):
                    # nếu là số nguyên thì định dạng có dấu chấm ngăn cách hàng nghìn
                if float(raw).is_integer():
                    display = f"{int(raw):,}".replace(",", ".")
                else:
                    display = f"{raw:,}".replace(",", ".")
            else:
                # trường hợp chuỗi (ví dụ Excel đã có "30/06/2025") -> giữ nguyên
                display = str(raw)
            row_dict[str(header)] = display
        rows.append(row_dict)
    return headers, rows

# --- Hàm thay placeholder trong 1 paragraph, cố gắng giữ nguyên định dạng run ---
def replace_placeholders_in_paragraph(paragraph, data_dict):
    # Quick replace when placeholder entirely inside a single run
    for key, val in data_dict.items():
        placeholder = "{{" + str(key) + "}}"
        # first pass: replace in any single run that contains the whole placeholder
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(val))

    # second pass: handle placeholders split across multiple runs
    # build full text and search
    full_text = "".join([r.text for r in paragraph.runs])
    for key, val in data_dict.items():
        placeholder = "{{" + str(key) + "}}"
        start_index = full_text.find(placeholder)
        # while there is an occurrence, handle it
        while start_index != -1:
            # find which runs cover the placeholder
            char_count = 0
            start_run_idx = None
            end_run_idx = None
            for i, r in enumerate(paragraph.runs):
                next_count = char_count + len(r.text)
                if start_run_idx is None and start_index < next_count:
                    start_run_idx = i
                    start_offset = start_index - char_count
                if start_index + len(placeholder) <= next_count:
                    end_run_idx = i
                    end_offset = start_index + len(placeholder) - char_count
                    break
                char_count = next_count

            if start_run_idx is None or end_run_idx is None:
                break  # shouldn't happen, but safe-guard

            # combine text of runs that cover the placeholder
            combined = "".join(r.text for r in paragraph.runs[start_run_idx:end_run_idx+1])
            # replace placeholder inside this combined chunk
            new_combined = combined.replace(placeholder, str(val), 1)

            # put new_combined into the first run, clear the other runs
            paragraph.runs[start_run_idx].text = new_combined
            for j in range(start_run_idx+1, end_run_idx+1):
                paragraph.runs[j].text = ""

            # rebuild full_text and search next occurrence
            full_text = "".join([r.text for r in paragraph.runs])
            start_index = full_text.find(placeholder, start_index + len(str(val)))

# --- Hàm thay placeholders trong toàn bộ document (paragraphs + tables) ---
def fill_template(template_path, output_path, data_dict):
    doc = Document(template_path)

    # paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, data_dict)

    # tables (duyệt mọi ô)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, data_dict)

    doc.save(output_path)


if __name__ == "__main__":
    # Khi chạy trực tiếp, dùng các giá trị mặc định cũ để giữ tương thích
    excel_file = "data2.xlsx"
    template_file = "template2.docx"
    output_folder = DEFAULT_OUTPUT_FOLDER

    print("Headers read from Excel:")
    headers, rows = read_excel_display_rows(excel_file)
    print(headers)

    created = generate_from_files(excel_file, template_file, output_folder)
    for p in created:
        print(f"Created: {p}")

    print("Done.")
